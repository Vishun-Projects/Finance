"""
Multi-Format File Parser
========================
Supports parsing financial data from multiple file formats:
- PDF (bank statements)
- XLS/XLSX (Excel files)
- DOC/DOCX (Word documents)
- TXT (text files)

Extracts transaction data and converts to standardized format.
"""

import re
import pandas as pd
import os
from pathlib import Path
from datetime import datetime
import chardet

# Excel support
try:
    import openpyxl
    from openpyxl import load_workbook
except ImportError:
    openpyxl = None

try:
    import xlrd
except ImportError:
    xlrd = None

# Word document support
try:
    from docx import Document
except ImportError:
    Document = None

# PDF support
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None


def detect_file_encoding(file_path):
    """Detect file encoding for text files."""
    with open(file_path, 'rb') as f:
        raw_data = f.read()
        result = chardet.detect(raw_data)
        return result['encoding']


def parse_pdf_file(file_path):
    """Parse PDF files (existing functionality)."""
    if not pdfplumber:
        raise ImportError("pdfplumber not available for PDF parsing")
    
    # Import the existing accurate parser
    import sys
    sys.path.append(os.path.dirname(__file__))
    from accurate_parser import parse_bank_statement_accurately
    
    return parse_bank_statement_accurately(file_path)


def parse_excel_file(file_path):
    """Parse Excel files (XLS/XLSX) with specific structure for bank statements."""
    try:
        # Read Excel file without headers to get raw data
        df = pd.read_excel(file_path, engine='openpyxl', header=None)
        
        if df.empty:
            raise Exception("Excel file is empty or has no readable data")
        
        transactions = []
        
        # Process each row looking for transaction data
        for index, row in df.iterrows():
            # Skip empty rows
            if row.isna().all():
                continue
            
            # Check if this row contains transaction data
            # Look for date pattern in column 1 (index 1)
            date_value = None
            if pd.notna(row.iloc[1]) and isinstance(row.iloc[1], str):
                date_match = re.match(r'(\d{1,2}\s+[A-Za-z]{3}\s+\d{4})', str(row.iloc[1]).strip())
                if date_match:
                    date_value = date_match.group(1)
            
            # If we found a date, this is likely a transaction row
            if date_value:
                # Get transaction details from column 2 (index 2)
                description = str(row.iloc[2]) if pd.notna(row.iloc[2]) else ''
                
                # Get debit amount from column 8 (index 8)
                debit_str = str(row.iloc[8]) if pd.notna(row.iloc[8]) else ''
                debit_amount = 0.0
                if 'INR' in debit_str and debit_str.strip() != ' - ':
                    debit_match = re.search(r'INR\s*([0-9,]+(?:\.[0-9]{2})?)', debit_str)
                    if debit_match:
                        debit_amount = float(debit_match.group(1).replace(',', ''))
                
                # Get credit amount from column 10 (index 10)
                credit_str = str(row.iloc[10]) if pd.notna(row.iloc[10]) else ''
                credit_amount = 0.0
                if 'INR' in credit_str and credit_str.strip() != ' - ':
                    credit_match = re.search(r'INR\s*([0-9,]+(?:\.[0-9]{2})?)', credit_str)
                    if credit_match:
                        credit_amount = float(credit_match.group(1).replace(',', ''))
                
                # Skip if no meaningful data found
                if not description or (debit_amount == 0 and credit_amount == 0):
                    continue
                
                # Extract store and commodity information
                store, commodity, clean_description = extract_store_and_commodity(description)
                
                # Determine transaction type
                if debit_amount > 0:
                    transaction_type = 'expense'
                    amount = debit_amount
                else:
                    transaction_type = 'income'
                    amount = credit_amount
                
                # Combine store, commodity, and remarks
                combined_remarks = []
                if store:
                    combined_remarks.append(store)
                if commodity:
                    combined_remarks.append(commodity)
                combined_remarks = ' | '.join(combined_remarks) if combined_remarks else ''
                
                transactions.append({
                    'date': date_value,
                    'description': clean_description,
                    'raw': description,
                    'remarks': combined_remarks,
                    'amount': amount,
                    'type': transaction_type,
                    'debit': debit_amount,
                    'credit': credit_amount,
                    'balance': None,
                    'page': 'Excel',
                    'line': str(index + 1),
                    'date_iso': pd.to_datetime(date_value).strftime('%Y-%m-%d'),
                    'store': store,
                    'commodity': commodity
                })
        
        return pd.DataFrame(transactions)
        
    except Exception as e:
        raise Exception(f"Failed to parse Excel file: {str(e)}")


def parse_word_file(file_path):
    """Parse Word documents (DOC/DOCX)."""
    if not Document:
        raise ImportError("python-docx not available for Word document parsing")
    
    try:
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        # Join all text
        full_text = '\n'.join(text_content)
        
        # Parse using similar logic as PDF parser
        return parse_text_content(full_text, source="Word Document")
        
    except Exception as e:
        raise Exception(f"Failed to parse Word document: {str(e)}")


def parse_text_file(file_path):
    """Parse text files (TXT)."""
    try:
        # Detect encoding
        encoding = detect_file_encoding(file_path)
        if not encoding:
            encoding = 'utf-8'
        
        # Read file content
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()
        
        return parse_text_content(content, source="Text File")
        
    except Exception as e:
        raise Exception(f"Failed to parse text file: {str(e)}")


def extract_store_and_commodity(description):
    """Extract store name and commodity from transaction description."""
    store = None
    commodity = None
    clean_description = description
    
    # Simple pattern: TRANSACTION_CODE/Store Name /rest_of_description
    # Examples from your data:
    # YESB0PTMUPI/Indian Railways Ticketing /XXXXX /paytm-64670120@paytm /UPI/411950138862/UPI/BRANCH : ATM SERVICE BRANCH
    # HDFC0002504/MAMTA MUNSHEELAL VISHWAKARMA /XXXXX06233/mamtavishwakarma0948@okhdfcbank /UPI/412484498384/UPI/BRANCH : ATM SERVICE BRANCH
    # KKBK0001432/AMARAWATIDEVI MUNSHEELAL VISHW /XXXXX88809/amravatidevivishwakarma61@okhdfcbank/UPI/415078935523/recharge /BRANCH : ATM SERVICE BRANCH
    
    # Extract store name - everything after the first slash until the next slash or end
    store_match = re.search(r'^[A-Z0-9_]+/([^/]+?)(?:\s*/\s*|$)', description)
    if store_match:
        store = store_match.group(1).strip()
        # Clean up store name - remove extra spaces
        store = re.sub(r'\s+', ' ', store).strip()
    
    # Extract commodity - look for meaningful words in various patterns
    commodity_patterns = [
        # Pattern 1: /UPI/numbers/commodity /BRANCH
        r'/UPI/\d+/([a-zA-Z][a-zA-Z\s]+?)(?:\s*/\s*BRANCH|$)',
        # Pattern 2: /Pay to something /BRANCH
        r'/Pay to ([a-zA-Z][a-zA-Z\s]+?)(?:\s*/\s*BRANCH|$)',
        # Pattern 3: /XXXXX/commodity /UPI
        r'/XXXXX[^/]*/([a-zA-Z][a-zA-Z\s]+?)(?:\s*/\s*UPI)',
        # Pattern 4: /commodity /BRANCH (at the end)
        r'/([a-zA-Z][a-zA-Z\s]+?)(?:\s*/\s*BRANCH\s*:\s*ATM SERVICE BRANCH)$',
    ]
    
    for pattern in commodity_patterns:
        commodity_match = re.search(pattern, description)
        if commodity_match:
            candidate = commodity_match.group(1).strip()
            # Skip technical codes and meaningless words
            if (candidate and len(candidate) > 1 and 
                candidate.lower() not in ['upi', 'branch', 'atm', 'service', 'pay', 'to', 'bharatpe', 'merc', 'yesbankltd', 'kotak', 'hdfcbank', 'okaxis']):
                commodity = candidate
                break
    
    # Clean up description - remove technical parts
    clean_description = description
    # Remove UPI IDs and technical codes
    clean_description = re.sub(r'/\s*[A-Z0-9@._-]+@[a-zA-Z0-9.-]+', '', clean_description)  # Remove email-like UPI IDs
    clean_description = re.sub(r'/\s*UPI/\d+', '', clean_description)  # Remove UPI/numbers
    clean_description = re.sub(r'/\s*BRANCH.*', '', clean_description)  # Remove branch info
    clean_description = re.sub(r'/\s*XXXXX[^/]*', '', clean_description)  # Remove XXXXX codes
    clean_description = re.sub(r'/\s*[A-Z0-9_]+', '', clean_description)  # Remove other codes
    clean_description = re.sub(r'\s+', ' ', clean_description).strip()  # Clean whitespace
    
    return store, commodity, clean_description

def parse_text_content(text, source="Unknown"):
    """Parse text content for transaction data."""
    # Normalize line breaks
    text = re.sub(r'\r\n|\r', '\n', text)
    text = re.sub(r'\n{2,}', '\n\n', text)
    
    transactions = []
    lines = text.split('\n')
    
    # Process structured text format (like our sample)
    current_transaction = {}
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        
        # Check for date line
        if line.startswith('Date:'):
            # Save previous transaction if exists
            if current_transaction and 'date' in current_transaction:
                transactions.append(current_transaction)
            
            # Start new transaction
            date_match = re.search(r'Date:\s*(\d{1,2}\s+[A-Za-z]{3}\s+\d{4})', line)
            if date_match:
                current_transaction = {
                    'date': date_match.group(1),
                    'description': '',
                    'remarks': None,
                    'debit': 0.0,
                    'credit': 0.0,
                    'balance': None,
                    'page': source,
                    'line': str(i + 1),
                    'date_iso': None,
                    'store': None,
                    'commodity': None
                }
        
        # Check for transaction line
        elif line.startswith('Transaction:'):
            if current_transaction:
                transaction_match = re.search(r'Transaction:\s*(.+)', line)
                if transaction_match:
                    raw_description = transaction_match.group(1).strip()
                    # Extract store and commodity information
                    store, commodity, clean_description = extract_store_and_commodity(raw_description)
                    current_transaction['description'] = clean_description
                    current_transaction['raw'] = raw_description
                    current_transaction['store'] = store
                    current_transaction['commodity'] = commodity
        
        # Check for amount line
        elif line.startswith('Amount:'):
            if current_transaction:
                amount_match = re.search(r'Amount:\s*INR\s*([0-9,]+(?:\.[0-9]{2})?)', line)
                if amount_match:
                    amount = float(amount_match.group(1).replace(',', ''))
                    current_transaction['amount'] = amount
        
        # Check for type line
        elif line.startswith('Type:'):
            if current_transaction and 'amount' in current_transaction:
                type_match = re.search(r'Type:\s*(Debit|Credit)', line)
                if type_match:
                    txn_type = type_match.group(1)
                    amount = current_transaction['amount']
                    
                    if txn_type == 'Debit':
                        current_transaction['debit'] = amount
                        current_transaction['credit'] = 0.0
                        current_transaction['type'] = 'expense'
                        current_transaction['amount'] = amount
                    else:
                        current_transaction['debit'] = 0.0
                        current_transaction['credit'] = amount
                        current_transaction['type'] = 'income'
                        current_transaction['amount'] = amount
                    
                    # Set date_iso
                    try:
                        current_transaction['date_iso'] = pd.to_datetime(current_transaction['date']).strftime('%Y-%m-%d')
                    except:
                        current_transaction['date_iso'] = None
                    
                    # Combine store, commodity, and remarks
                    combined_remarks = []
                    if current_transaction.get('store'):
                        combined_remarks.append(current_transaction['store'])
                    if current_transaction.get('commodity'):
                        combined_remarks.append(current_transaction['commodity'])
                    current_transaction['remarks'] = ' | '.join(combined_remarks) if combined_remarks else ''
                    
                    # Remove temporary amount field
                    del current_transaction['amount']
    
    # Add the last transaction
    if current_transaction and 'date' in current_transaction:
        transactions.append(current_transaction)
    
    # If no structured format found, try generic parsing
    if not transactions:
        # Common date patterns
        date_patterns = [
            r'\d{1,2}\s+[A-Za-z]{3}\s+\d{4}',  # 08 Sep 2025
            r'\d{1,2}/\d{1,2}/\d{4}',          # 08/09/2025
            r'\d{1,2}-\d{1,2}-\d{4}',          # 08-09-2025
            r'\d{4}-\d{1,2}-\d{1,2}',          # 2025-09-08
        ]
        
        # Amount patterns
        amount_patterns = [
            r'INR\s*([0-9,]+(?:\.[0-9]{2})?)',
            r'₹\s*([0-9,]+(?:\.[0-9]{2})?)',
            r'\$\s*([0-9,]+(?:\.[0-9]{2})?)',
            r'([0-9,]+(?:\.[0-9]{2})?)\s*(?:INR|₹|\$)',
            r'([+-]?\s*[0-9,]+(?:\.[0-9]{2})?)',
        ]
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Look for date patterns
            date_match = None
            for pattern in date_patterns:
                match = re.search(pattern, line)
                if match:
                    date_match = match.group(0)
                    break
            
            if not date_match:
                continue
            
            # Look for amount patterns
            amounts = []
            for pattern in amount_patterns:
                matches = re.finditer(pattern, line)
                for match in matches:
                    try:
                        amount_str = match.group(1).replace(',', '')
                        amount = float(amount_str)
                        amounts.append((amount, match.start()))
                    except:
                        continue
            
            if not amounts:
                continue
            
            # Determine debit/credit
            debit_amount = 0.0
            credit_amount = 0.0
            
            # Look for transaction amount (usually the first or second amount)
            if len(amounts) >= 1:
                txn_amount, txn_pos = amounts[0]
                
                # Check if amount is negative
                if txn_amount < 0:
                    debit_amount = abs(txn_amount)
                else:
                    credit_amount = txn_amount
            
            # Clean description
            description = line
            description = re.sub(re.escape(date_match), '', description, count=1).strip()
            
            # Remove amount patterns from description
            for pattern in amount_patterns:
                description = re.sub(pattern, '', description)
            
            description = re.sub(r'\s{2,}', ' ', description).strip()
            
            # Extract remarks
            remark_match = re.search(r'/([A-Za-z][A-Za-z0-9 _.-]{2,})$', description)
            remark = None
            if remark_match:
                remark = remark_match.group(1).strip()
                description = description[:remark_match.start()].strip(" /")
            
            if debit_amount > 0 or credit_amount > 0:
                # Extract store and commodity information
                store, commodity, clean_description = extract_store_and_commodity(description)
                
                # Determine transaction type and amount
                if debit_amount > 0:
                    transaction_type = 'expense'
                    amount = debit_amount
                else:
                    transaction_type = 'income'
                    amount = credit_amount
                
                # Combine store, commodity, and existing remarks
                combined_remarks = []
                if store:
                    combined_remarks.append(store)
                if commodity:
                    combined_remarks.append(commodity)
                if remark:
                    combined_remarks.append(remark)
                combined_remarks = ' | '.join(combined_remarks) if combined_remarks else ''
                    
                transactions.append({
                    'date': date_match,
                    'description': clean_description,
                    'raw': original_line if 'original_line' in locals() else description_line,
                    'remarks': combined_remarks,
                    'amount': amount,
                    'type': transaction_type,
                    'debit': debit_amount,
                    'credit': credit_amount,
                    'balance': None,
                    'page': source,
                    'line': str(i + 1),
                    'date_iso': pd.to_datetime(date_match, errors='coerce').strftime('%Y-%m-%d') if pd.to_datetime(date_match, errors='coerce') is not pd.NaT else None,
                    'store': store,
                    'commodity': commodity
                })
    
    return pd.DataFrame(transactions)


def parse_file(file_path):
    """Main function to parse any supported file format."""
    file_path = Path(file_path)
    
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_extension = file_path.suffix.lower()
    
    print(f"Parsing {file_extension.upper()} file: {file_path.name}")
    
    try:
        if file_extension == '.pdf':
            return parse_pdf_file(file_path)
        elif file_extension in ['.xls', '.xlsx']:
            return parse_excel_file(file_path)
        elif file_extension in ['.doc', '.docx']:
            return parse_word_file(file_path)
        elif file_extension == '.txt':
            return parse_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
            
    except Exception as e:
        raise Exception(f"Failed to parse {file_extension.upper()} file: {str(e)}")


def main():
    """Test function for command line usage."""
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python multi_format_parser.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    try:
        df = parse_file(file_path)
        print(f"SUCCESS: Extracted {len(df)} transactions from {file_path}")
        
        # Save to CSV
        output_file = f"parsed_{Path(file_path).stem}.csv"
        df.to_csv(output_file, index=False)
        print(f"Results saved to: {output_file}")
        
    except Exception as e:
        print(f"ERROR: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
